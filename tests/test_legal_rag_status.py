# tests/test_legal_rag_status.py
"""
[FIX 2026-08-09 — LLM lỗi phải BÁO RÕ, không giấu thành data thật] Test trạng thái
truy xuất căn cứ pháp lý (legal_rag_status) của RAG Agent.

Vấn đề gốc: khi LLM API lỗi (timeout/rate limit/sai key) hoặc chưa cấu hình
LLM_API_KEY, regulation_rag.py trả mock content và alert_report.py in thẳng vào
STR như căn cứ pháp lý thật — nguy hiểm vì STR là tài liệu nộp cơ quan chức năng.

Fix:
- regulation_rag.py trả legal_citations = [] + legal_rag_status = "UNAVAILABLE"
  khi lỗi (CẢ 2 trường hợp: MOCK_NO_KEY thiếu key và UNAVAILABLE API lỗi) —
  KHÔNG đẩy mock content vào legal_citations.
- alert_report.py khi UNAVAILABLE in cảnh báo tường minh vào STR, không để trống
  và không in nội dung không rõ nguồn gốc.
- legal_rag_status KHÔNG ảnh hưởng Decision Engine (REPORT/REVIEW/PASS quyết
  định độc lập trước khi RAG chạy).

Test cover:
1. API lỗi (UNAVAILABLE) → state có legal_rag_status="UNAVAILABLE",
   legal_citations=[] (không chứa mock), STR có cảnh báo + không chứa mock text.
2. Thiếu key (MOCK_NO_KEY) → xử lý GIỐNG HỆT UNAVAILABLE (cùng cảnh báo STR,
   không in mock).
3. Trường hợp OK → legal_rag_status="OK", citations vẫn được in vào STR (không
   bị ảnh hưởng bởi fix).
4. call_llm_api (HỢP ĐỒNG API với api/main.py chat) vẫn trả str như cũ.
"""
import os

import pytest

from core.state import AMLState
from core.config import REPORTS_OUTPUT_DIR
from agents.alert_report import generate_alert_report
from agents import regulation_rag
from agents.regulation_rag import (
    call_llm_api,
    run_regulation_rag,
)

# =============================================================================
# Helpers
# =============================================================================


def _read_docx_text(path: str) -> str:
    """Đọc toàn bộ paragraph + table text từ file .docx để assert."""
    from docx import Document

    doc = Document(path)
    parts = []
    for para in doc.paragraphs:
        parts.append(para.text)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                parts.append(cell.text)
    return "\n".join(parts)


def _rag_state(tx_hash: str) -> AMLState:
    """State tối thiểu đủ cho run_regulation_rag chạy (không cần ChromaDB thật)."""
    return AMLState(
        tx_hash=tx_hash,
        wallet_from="0x28C6c06298d514Db089934071355E5743bf21d60",
        wallet_to="0x000000000000000000000000000000000000dEaD",
        amount_vnd=200_000_000.0,
        hashed_fullname="d39fb3f8f4b0a9c8b1f0d8f5b9a0e1c2d3e4f5a6b7c8d9e0f1a2b3c4d5e6f7a8",
        hashed_id_number="a591a6d40bf42040",
        hashed_account_number="2c26b46b68ffc68f",
        classifier_score=0.0,
        graph_score=0.0,
        decision="REPORT",
        decision_reason="Test legal rag status.",
        decision_evidence=["Test trigger"],
        sanction_result={"is_match": False},
        graph_analysis_status="NO_GRAPH_DATA",
        graph_data_available=False,
        sanction_path_found=None,
        insufficient_data=False,
    )


def _report_state(tx_hash: str, **extra) -> AMLState:
    """State đã qua decision=REPORT (case_status=pending_review) cho generate_alert_report."""
    state = AMLState(
        tx_hash=tx_hash,
        wallet_from="0x28C6c06298d514Db089934071355E5743bf21d60",
        wallet_to="0x000000000000000000000000000000000000dEaD",
        amount_vnd=200_000_000.0,
        hashed_fullname="d39fb3f8f4b0a9c8b1f0d8f5b9a0e1c2d3e4f5a6b7c8d9e0f1a2b3c4d5e6f7a8",
        hashed_id_number="a591a6d40bf42040",
        hashed_account_number="2c26b46b68ffc68f",
        classifier_score=0.0,
        graph_score=0.0,
        case_status="pending_review",
        decision="REPORT",
        decision_reason="Test legal rag status.",
        decision_evidence=["Test trigger"],
        approval_status="pending",
        sanction_result={"is_match": False},
        aggregation_status="not_assessed",
        structuring_detected=None,
        graph_analysis_status="NO_GRAPH_DATA",
        graph_data_available=False,
        sanction_path_found=None,
        insufficient_data=False,
    )
    state.update(extra)
    return state


@pytest.fixture(autouse=True)
def _no_chromadb():
    """
    Cô lập test khỏi ChromaDB thật: query_legal_docs trả rỗng ({} là kết quả
    hợp lệ — _merge_unique lọc rỗng). KHÔNG cần vector DB cho các assert dưới đây.
    """
    original = regulation_rag.query_legal_docs
    regulation_rag.query_legal_docs = lambda query_text, n_results: {}
    yield
    regulation_rag.query_legal_docs = original


# =============================================================================
# 1. API lỗi (UNAVAILABLE) → STR có cảnh báo, KHÔNG có mock như thật
# =============================================================================

def test_api_error_marks_unavailable_and_no_mock_into_citations(monkeypatch):
    """LLM API raise/trả lỗi → legal_citations=[] + legal_rag_status=UNAVAILABLE + STR cảnh báo."""
    mock_error_text = (
        "[LLM API LỖI - dùng tạm Mock Output]: Không gọi được Gemini API "
        "(mock timeout). Giao dịch vi phạm quy định báo cáo (Thông tư 27)..."
    )
    monkeypatch.setattr(
        regulation_rag,
        "_call_llm_api_internal",
        lambda prompt: (mock_error_text, "UNAVAILABLE"),
    )

    state = run_regulation_rag(_rag_state("tx_rag_api_error"))
    assert state["legal_rag_status"] == "UNAVAILABLE"
    assert state["legal_citations"] == []
    assert state["legal_rag_error"] is not None
    # KHÔNG bao giờ mock content lọt vào field citations trông như data thật:
    assert "[Mock LLM Output]" not in str(state["legal_citations"])
    assert "[LLM API LỖI" not in str(state["legal_citations"])

    # STR phải có cảnh báo tường minh + KHÔNG chứa mock content.
    report_state = _report_state("tx_rag_api_error", legal_rag_status=state["legal_rag_status"], legal_rag_error=state["legal_rag_error"])
    updated = generate_alert_report(report_state)
    text = _read_docx_text(updated["report_path"])

    assert "Không thể truy xuất căn cứ pháp lý tự động" in text
    assert "Chuyên viên AML phải tự tra cứu" in text
    assert "[Mock LLM Output]" not in text
    assert "[LLM API LỖI" not in text


# =============================================================================
# 2. Thiếu key (MOCK_NO_KEY) → xử lý GIỐNG HỆT UNAVAILABLE
# =============================================================================

def test_missing_key_treated_same_as_unavailable(monkeypatch):
    """
    Feedback duyệt: KHÔNG có lý do 'thiếu key' an toàn hơn 'API lỗi' — MOCK_NO_KEY
    phải dẫn tới CÙNG cảnh báo trong STR và không in mock vào STR.
    """
    mock_no_key_text = (
        "[Mock LLM Output]: Dựa trên truy vấn, giao dịch vi phạm quy định báo cáo "
        "(Thông tư 27) hoặc phát sinh nghĩa vụ thuế (Thông tư 32)... "
        "(Vui lòng điền LLM_API_KEY vào .env để chạy thật)"
    )
    monkeypatch.setattr(
        regulation_rag,
        "_call_llm_api_internal",
        lambda prompt: (mock_no_key_text, "MOCK_NO_KEY"),
    )

    state = run_regulation_rag(_rag_state("tx_rag_no_key"))
    assert state["legal_rag_status"] == "UNAVAILABLE"  # Cùng status — KHÔNG phân biệt an toàn hơn
    assert state["legal_citations"] == []
    assert state["legal_rag_error"] is not None
    assert "LLM_API_KEY" in state["legal_rag_error"]

    report_state = _report_state("tx_rag_no_key", legal_rag_status=state["legal_rag_status"], legal_rag_error=state["legal_rag_error"])
    updated = generate_alert_report(report_state)
    text = _read_docx_text(updated["report_path"])

    assert "Không thể truy xuất căn cứ pháp lý tự động" in text
    assert "Chuyên viên AML phải tự tra cứu" in text
    assert "[Mock LLM Output]" not in text


# =============================================================================
# 3. Trường hợp OK → không bị ảnh hưởng, citations vẫn in vào STR
# =============================================================================

def test_ok_status_keeps_real_citations_in_str(monkeypatch):
    """LLM hoạt động tốt → legal_rag_status=OK, citations thật vẫn vào STR."""
    ok_json = (
        '{"legal_citations": ['
        '  {"source_file": "thong_tu_27_2025.txt", "dieu_khoan": "Điều 7", '
        '   "noi_dung_tom_tat": "Nghĩa vụ báo cáo giao dịch đáng ngờ.", '
        '   "ly_do_ap_dung": "Giao dịch có dấu hiệu đáng ngờ."}'
        "]}"
    )
    monkeypatch.setattr(
        regulation_rag,
        "_call_llm_api_internal",
        lambda prompt: (ok_json, "OK"),
    )

    state = run_regulation_rag(_rag_state("tx_rag_ok"))
    assert state["legal_rag_status"] == "OK"
    assert state["legal_rag_error"] is None
    assert state["legal_citations"], "Phải có citations khi LLM OK"
    # Nguồn được chuẩn hoá từ source_file qua CANONICAL_SOURCE_NAMES.
    assert state["legal_citations"][0]["source"] == "Thông tư 27/2025/TT-NHNN"
    assert state["legal_citations"][0]["dieu_khoan"] == "Điều 7"

    report_state = _report_state("tx_rag_ok", legal_rag_status=state["legal_rag_status"], legal_rag_error=state["legal_rag_error"], legal_citations=state["legal_citations"])
    updated = generate_alert_report(report_state)
    text = _read_docx_text(updated["report_path"])

    # Citations THẬT vẫn được in vào STR bình thường — không bị fix làm mất.
    assert "Thông tư 27/2025/TT-NHNN — Điều 7" in text
    assert "Giao dịch có dấu hiệu đáng ngờ." in text
    # KHÔNG có cảnh báo UNAVAILABLE trong case OK.
    assert "Không thể truy xuất căn cứ pháp lý tự động" not in text


# =============================================================================
# 4. Hợp đồng API call_llm_api (api/main.py chat) KHÔNG vỡ
# =============================================================================

def test_call_llm_api_contract_unchanged(monkeypatch):
    """call_llm_api vẫn trả str thuần (api/main.py chat phụ thuộc)."""
    raw = "chat answer text"
    monkeypatch.setattr(
        regulation_rag,
        "_call_llm_api_internal",
        lambda prompt: (raw, "OK"),
    )
    result = call_llm_api("prompt")
    assert isinstance(result, str)
    assert result == raw
