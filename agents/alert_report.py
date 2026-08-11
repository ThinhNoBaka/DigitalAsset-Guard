# agents/alert_report.py

"""
Report Assistant - Tổng hợp báo cáo STR (Mẫu 04) khi
case_status == "pending_review".

IMPORTANT:
- Report Assistant KHÔNG tự quyết định có lập STR hay không.
- Quyết định đã được các agent trước đó ghi vào AMLState.
- Report Assistant chỉ tổng hợp và định dạng dữ liệu thành DOCX.
- risk_assessment_score là điểm rủi ro NỘI BỘ của DigitalAsset Guard,
  không phải ngưỡng pháp lý.
- Phần phân tích XAI/Graph được đưa vào phụ lục nội bộ để hỗ trợ
  AML Officer, không thay thế nội dung của biểu mẫu STR.
"""

import os
from datetime import datetime
from typing import Any

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import (
    WD_TABLE_ALIGNMENT,
    WD_CELL_VERTICAL_ALIGNMENT,
)

from core.privacy_layer import assert_no_raw_pii
from core.state import AMLState
from core.config import REPORTS_OUTPUT_DIR


# ============================================================
# CONFIGURATION
# ============================================================

os.makedirs(
    REPORTS_OUTPUT_DIR,
    exist_ok=True,
)


# ============================================================
# DOCUMENT HELPERS
# ============================================================

def _set_run_font(run, size=12, bold=False, italic=False):
    """Định dạng font thống nhất cho văn bản."""
    run.font.name = "Times New Roman"
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic


def _add_centered_line(
    doc: Document,
    text: str,
    bold: bool = False,
    italic: bool = False,
    size: int = 12,
):
    """Thêm một paragraph căn giữa."""
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    run = paragraph.add_run(text)
    _set_run_font(
        run,
        size=size,
        bold=bold,
        italic=italic,
    )

    return paragraph


def _add_field(
    doc: Document,
    label: str,
    value: Any = None,
):
    """
    Thêm một trường dạng:

    Label: Value
    """

    if value is None or value == "":
        value = "[CHƯA CUNG CẤP]"

    paragraph = doc.add_paragraph()

    label_run = paragraph.add_run(
        f"{label}: "
    )

    _set_run_font(
        label_run,
        bold=True,
    )

    value_run = paragraph.add_run(
        str(value)
    )

    _set_run_font(
        value_run
    )

    return paragraph


def _add_section_title(
    doc: Document,
    title: str,
):
    """Thêm tiêu đề section."""
    paragraph = doc.add_paragraph()

    run = paragraph.add_run(title)

    _set_run_font(
        run,
        size=12,
        bold=True,
    )

    return paragraph


def _add_table(
    doc: Document,
    headers,
    rows,
):
    """
    Tạo bảng DOCX đơn giản, phù hợp với report.
    """

    table = doc.add_table(
        rows=1,
        cols=len(headers),
    )

    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header
    header_cells = table.rows[0].cells

    for i, header in enumerate(headers):

        cell = header_cells[i]

        cell.text = str(header)

        cell.vertical_alignment = (
            WD_CELL_VERTICAL_ALIGNMENT.CENTER
        )

        for paragraph in cell.paragraphs:

            for run in paragraph.runs:
                _set_run_font(
                    run,
                    size=11,
                    bold=True,
                )

    # Rows
    for row in rows:

        cells = table.add_row().cells

        for i, value in enumerate(row):

            if value is None or value == "":
                value = "Không cung cấp"

            cells[i].text = str(value)

            cells[i].vertical_alignment = (
                WD_CELL_VERTICAL_ALIGNMENT.CENTER
            )

            for paragraph in cells[i].paragraphs:

                for run in paragraph.runs:
                    _set_run_font(
                        run,
                        size=11,
                    )

    return table


def _add_bullet(
    doc: Document,
    text: str,
):
    """Thêm bullet point."""
    paragraph = doc.add_paragraph(
        style="List Bullet"
    )

    run = paragraph.add_run(
        str(text)
    )

    _set_run_font(run)

    return paragraph


# ============================================================
# LEGAL CITATION HANDLER
# ============================================================

def _add_legal_citations(
    doc: Document,
    legal_citations,
):
    """
    Hỗ trợ cả 2 schema:

    1. List[str]

    2. List[dict]

    Ví dụ dict:

    {
        "source": "Thông tư 27/2025/TT-NHNN",
        "dieu_khoan": "Điều 6",
        "noi_dung_tom_tat": "...",
        "ly_do_ap_dung": "..."
    }
    """

    if not legal_citations:

        doc.add_paragraph(
            "Chưa có căn cứ pháp lý được RAG Agent "
            "cung cấp."
        )

        return

    for citation in legal_citations:

        # ----------------------------------------------------
        # Schema cũ: string
        # ----------------------------------------------------

        if isinstance(citation, str):

            _add_bullet(
                doc,
                citation,
            )

            continue

        # ----------------------------------------------------
        # Schema mới: dict
        # ----------------------------------------------------

        if isinstance(citation, dict):

            # Cho phép raw_text để backward-compatible
            if (
                "raw_text" in citation
                and len(citation) == 1
            ):
                _add_bullet(
                    doc,
                    citation["raw_text"],
                )

                continue

            source = citation.get(
                "source",
                "Không rõ nguồn",
            )

            article = citation.get(
                "dieu_khoan",
                "",
            )

            summary = citation.get(
                "noi_dung_tom_tat",
                "",
            )

            reason = citation.get(
                "ly_do_ap_dung",
                "",
            )

            title = source

            if article:
                title += f" — {article}"

            paragraph = doc.add_paragraph()

            run = paragraph.add_run(
                title
            )

            _set_run_font(
                run,
                bold=True,
            )

            if summary:

                _add_field(
                    doc,
                    "Nội dung",
                    summary,
                )

            if reason:

                _add_field(
                    doc,
                    "Lý do áp dụng",
                    reason,
                )

            continue

        # ----------------------------------------------------
        # Unknown schema
        # ----------------------------------------------------

        _add_bullet(
            doc,
            str(citation),
        )


# ============================================================
# SUSPICIOUS BEHAVIOR DESCRIPTION
# ============================================================

def _build_suspicious_description(
    state: AMLState,
) -> str:
    """
    Tổng hợp mô tả giao dịch đáng ngờ từ những gì
    đã có trong AMLState.

    Không tự suy luận thêm kết luận pháp lý.
    """

    tx_hash = state.get(
        "tx_hash",
        "UNKNOWN",
    )

    wallet_from = state.get(
        "wallet_from",
        "",
    )

    wallet_to = state.get(
        "wallet_to",
        "",
    )

    amount_vnd = (
        state.get(
            "amount_vnd",
            0,
        )
        or 0
    )

    decision_reason = state.get(
        "decision_reason",
        "",
    )

    parts = []

    # Transaction
    parts.append(
        f"Giao dịch {tx_hash}"
        f" có giá trị quy đổi "
        f"{amount_vnd:,.0f} VND."
    )

    if wallet_from:

        parts.append(
            f"Ví nguồn: {wallet_from}."
        )

    if wallet_to:

        parts.append(
            f"Ví đích: {wallet_to}."
        )

    # Decision reason
    if decision_reason:

        parts.append(
            f"Lý do được hệ thống ghi nhận: "
            f"{decision_reason}."
        )

    # Sanctions
    sanction_result = state.get(
        "sanction_result",
        {},
    ) or {}

    if sanction_result.get("is_match"):

        matched_wallet = sanction_result.get(
            "matched_wallet",
            "N/A",
        )

        program = sanction_result.get(
            "program",
            "N/A",
        )

        parts.append(
            "Kết quả sàng lọc sanctions ghi nhận "
            f"match với thực thể/ví {matched_wallet}, "
            f"chương trình {program}."
        )

    # Name fuzzy match
    if state.get(
        "name_similarity_warning",
        False,
    ):

        similarity = state.get(
            "name_similarity_score"
        )

        if similarity is not None:

            parts.append(
                "Hệ thống ghi nhận cảnh báo tương đồng "
                f"tên với mức {similarity:.2f}%; "
                "đây là tín hiệu fuzzy matching và "
                "cần được chuyên viên xác minh."
            )

        else:

            parts.append(
                "Hệ thống ghi nhận cảnh báo tương đồng "
                "tên; cần được chuyên viên xác minh."
            )

    # Graph
    community_id = state.get(
        "community_id"
    )

    suspicious_path = state.get(
        "suspicious_path",
        [],
    ) or []

    if community_id is not None:

        parts.append(
            f"Phân tích đồ thị xác định giao dịch "
            f"liên quan đến cộng đồng Louvain "
            f"{community_id}."
        )

    if suspicious_path:

        parts.append(
            "Đường đi đáng ngờ được ghi nhận: "
            + " → ".join(
                map(
                    str,
                    suspicious_path,
                )
            )
            + "."
        )

    # Decision evidence
    decision_evidence = state.get(
        "decision_evidence",
        [],
    ) or []

    if decision_evidence:

        parts.append(
            "Các bằng chứng được hệ thống sử dụng "
            "trong quá trình đánh giá gồm: "
            + "; ".join(
                map(
                    str,
                    decision_evidence,
                )
            )
            + "."
        )

    # [FIX 2026-08-08 — AUDIT ZERO-TX WALLET] Khi case_status thành pending_review
    # do THIẾU DỮ LIỆU on-chain (không phải do rủi ro cao thật sự), phải nêu rõ
    # để chuyên viên không hiểu nhầm mức độ nghiêm trọng. Decision Engine đã ghi
    # decision_evidence = "THIẾU DỮ LIỆU: ..." ở trên, nhưng bổ sung thêm 1 câu
    # khẳng định rõ ràng ngay trong mô tả giao dịch đáng ngờ.
    if state.get("insufficient_data", False):

        parts.append(
            "LƯU Ý QUAN TRỌNG: hệ thống chưa tìm thấy "
            "bất kỳ giao dịch on-chain nào của ví (0 giao "
            "dịch ETH và ERC20 trên Etherscan). Case này "
            "được chuyển sang xem xét do THIẾU DỮ LIỆU "
            "để đánh giá rủi ro, KHÔNG phải vì phát hiện "
            "dấu hiệu rủi ro cao — cần chuyên viên xác "
            "minh thủ công, không dùng classifier_score "
            "của hệ thống làm căn cứ."
        )

    # [FIX 2026-08-08 — OFF-CHAIN DATA SEPARATION] Khi Aggregation Monitor
    # KHÔNG có dữ liệu off-chain Core Banking (aggregation_status ==
    # "not_assessed"), phải nêu rõ structuring CHƯA được đánh giá — không
    # hiển thị như thể đã kiểm tra và pass (structuring_detected là None,
    # không phải False).
    if state.get("aggregation_status") == "not_assessed":

        parts.append(
            "Đánh giá structuring: Không đánh giá được "
            "(thiếu dữ liệu off-chain Core Banking). "
            "Không đồng nghĩa với việc đã kiểm tra và "
            "không phát hiện."
        )

    # Important disclaimer
    parts.append(
        "Các thông tin trên là kết quả hỗ trợ phân tích "
        "của DigitalAsset Guard và cần được chuyên viên "
        "AML kiểm tra, xác minh trước khi đưa ra quyết định "
        "cuối cùng."
    )

    return " ".join(parts)


# ============================================================
# MAIN REPORT GENERATOR
# ============================================================

def generate_alert_report(
    state: AMLState,
    *,
    report_customer_info=None,
    report_entity_info=None,
) -> AMLState:
    """
    Tạo báo cáo STR khi case_status == "pending_review".

    Report Assistant KHÔNG:
    - tự tính lại risk score;
    - tự thay đổi decision;
    - tự quyết định STR;
    - sử dụng 0.7 làm ngưỡng pháp lý.

    Report Assistant CHỈ:
    - đọc dữ liệu từ AMLState;
    - tạo DOCX;
    - lưu report_path;
    - đặt approval_status = pending để HITL xử lý.

    [FIX 2026-08-09 — STR thiếu thông tin định danh khách hàng]
    `report_customer_info` (PII plaintext: customer_name/customer_id_number/
    customer_account_number) và `report_entity_info` (config tĩnh hệ thống:
    reporting_entity_name/aml_responsible_person/reporter_name...) được truyền
    qua closure từ core/graph_builder.py::PipelineRun — KHÔNG nằm trong state
    channels, nên không bao giờ bị LangGraph checkpoint lưu (kể cả khi dùng
    checkpointer persistent), không bao giờ chạm assert_no_raw_pii, và không
    lẫn PII giữa các request (mỗi PipelineRun gọi build_pipeline mới).

    Nếu KHÔNG truyền (default None — vd. demo_runner gọi trực tiếp mà chưa
    chuyển qua PipelineRun), report vẫn in placeholder "[CHƯA CUNG CẤP]" /
    "[CHƯA CẤU HÌNH]" như trước — KHÔNG bao giờ fallback sang PII của request khác.
    """

    # ========================================================
    # 1. PRIVACY GATE
    # ========================================================

    assert_no_raw_pii(state)

    # ========================================================
    # 2. CASE STATUS GATE
    # ========================================================

    case_status = state.get(
        "case_status"
    )

    if case_status != "pending_review":

        state["report_path"] = None

        state["thought"] = (
            "ReportAssistant: Case không ở trạng thái "
            "pending_review, không tạo STR."
        )

        return state

    # ========================================================
    # 3. READ STATE
    # ========================================================

    tx_hash = state.get(
        "tx_hash",
        "UNKNOWN",
    )

    wallet_from = state.get(
        "wallet_from",
        "",
    )

    wallet_to = state.get(
        "wallet_to",
        "",
    )

    amount_vnd = (
        state.get(
            "amount_vnd",
            0,
        )
        or 0
    )

    classifier_score = (
        state.get(
            "classifier_score",
            0,
        )
        or 0
    )

    graph_score = (
        state.get(
            "graph_score",
            0,
        )
        or 0
    )

    # risk_assessment_score không còn được dùng để hiển thị (xem PHỤ LỤC A
    # bên dưới) — decision_engine.py đã bỏ tính điểm tổng hợp này khi
    # chuyển sang rule-based composite. Vẫn đọc ở đây (không xoá hẳn) để
    # nếu code cũ/khác còn ghi field này vào state thì không silently mất
    # thông tin — nhưng KHÔNG render vào báo cáo nữa vì dễ hiểu nhầm.
    risk_assessment_score = (
        state.get(
            "risk_assessment_score",
            0,
        )
        or 0
    )

    decision = state.get(
        "decision",
        "REPORT",
    )

    decision_reason = state.get(
        "decision_reason",
        "",
    )

    decision_evidence = (
        state.get(
            "decision_evidence",
            [],
        )
        or []
    )

    sanction_result = (
        state.get(
            "sanction_result",
            {},
        )
        or {}
    )

    name_similarity_warning = state.get(
        "name_similarity_warning",
        False,
    )

    name_similarity_score = state.get(
        "name_similarity_score"
    )

    name_similarity_matched_name = state.get(
        "name_similarity_matched_name"
    )

    legal_citations = (
        state.get(
            "legal_citations",
            [],
        )
        or []
    )

    # [FIX 2026-08-09 — LLM lỗi phải BÁO RÕ trong STR, không giấu thành data thật]
    # legal_rag_status do agents/regulation_rag.py ghi: "OK" (trích dẫn pháp lý
    # THẬT) hoặc "UNAVAILABLE" (không truy xuất được — thiếu LLM_API_KEY hoặc
    # API call lỗi; khi đó legal_citations LUÔN = [], KHÔNG bao giờ chứa mock).
    # KHÔNG ảnh hưởng Decision Engine — REPORT/REVIEW/PASS đã được quyết định
    # trước khi RAG chạy.
    legal_rag_status = state.get(
        "legal_rag_status"
    )

    legal_rag_error = state.get(
        "legal_rag_error"
    )

    community_id = state.get(
        "community_id"
    )

    suspicious_path = (
        state.get(
            "suspicious_path",
            [],
        )
        or []
    )

    thought = state.get(
        "thought",
        "",
    )

    # [FIX 2026-08-09 — STR thiếu thông tin định danh khách hàng]
    # Đọc PII plaintext + config entity từ tham số truyền qua closure (KHÔNG
    # nằm trong state channels — xem docstring generate_alert_report).
    # Nếu None (default — demo_runner gọi trực tiếp chưa qua PipelineRun),
    # fallback rỗng → các field bên dưới in "[CHƯA CUNG CẤP]"/"[CHƯA CẤU HÌNH]"
    # như trước, KHÔNG bao giờ lẫn PII của request khác.
    rc = report_customer_info or {}
    re = report_entity_info or {}

    # ========================================================
    # 4. CREATE DOCUMENT
    # ========================================================

    doc = Document()

    # Normal style
    normal_style = doc.styles["Normal"]

    normal_style.font.name = (
        "Times New Roman"
    )

    normal_style.font.size = Pt(12)

    # ========================================================
    # 5. HEADER
    # ========================================================

    _add_centered_line(
        doc,
        "CỘNG HÒA XÃ HỘI CHỦ NGHĨA VIỆT NAM",
        bold=True,
        size=13,
    )

    _add_centered_line(
        doc,
        "Độc lập - Tự do - Hạnh phúc",
        bold=True,
        size=13,
    )

    _add_centered_line(
        doc,
        "---------------***---------------",
    )

    doc.add_paragraph()

    _add_centered_line(
        doc,
        "BÁO CÁO GIAO DỊCH ĐÁNG NGỜ",
        bold=True,
        size=15,
    )

    _add_centered_line(
        doc,
        "Mẫu số 04 - Ban hành kèm theo "
        "Thông tư số 27/2025/TT-NHNN",
        italic=True,
    )

    _add_centered_line(
        doc,
        "Ngày lập báo cáo: "
        + datetime.now().strftime("%d/%m/%Y %H:%M:%S"),
    )

    doc.add_paragraph()

    # ========================================================
    # 6. PHẦN I
    # ========================================================

    _add_section_title(
        doc,
        "PHẦN I. THÔNG TIN VỀ ĐỐI TƯỢNG BÁO CÁO",
    )

    _add_field(
        doc,
        "Tên đối tượng báo cáo",
        re.get(
            "reporting_entity_name",
            "[CHƯA CẤU HÌNH]",
        ),
    )

    _add_field(
        doc,
        "Mã đối tượng báo cáo",
        re.get(
            "reporting_entity_code",
            "[CHƯA CẤU HÌNH]",
        ),
    )

    _add_field(
        doc,
        "Địa chỉ",
        re.get(
            "reporting_entity_address",
            "[CHƯA CẤU HÌNH]",
        ),
    )

    _add_field(
        doc,
        "Điện thoại",
        re.get(
            "reporting_entity_phone",
            "[CHƯA CẤU HÌNH]",
        ),
    )

    _add_field(
        doc,
        "Email",
        re.get(
            "reporting_entity_email",
            "[CHƯA CẤU HÌNH]",
        ),
    )

    _add_field(
        doc,
        "Người chịu trách nhiệm về phòng, "
        "chống rửa tiền",
        re.get(
            "aml_responsible_person",
            "[CHƯA CẤU HÌNH]",
        ),
    )

    _add_field(
        doc,
        "Chức vụ",
        re.get(
            "aml_responsible_position",
            "[CHƯA CẤU HÌNH]",
        ),
    )

    _add_field(
        doc,
        "Người lập báo cáo",
        re.get(
            "reporter_name",
            "[CHƯA CẤU HÌNH]",
        ),
    )

    # ========================================================
    # 7. PHẦN II
    # ========================================================

    _add_section_title(
        doc,
        "PHẦN II. THÔNG TIN VỀ CÁ NHÂN, "
        "TỔ CHỨC THỰC HIỆN GIAO DỊCH ĐÁNG NGỜ",
    )

    # LƯU Ý:
    # Các thông tin định danh đã được Privacy Layer kiểm tra.
    # Report Assistant chỉ đọc field đã được phép sử dụng.

    _add_field(
        doc,
        "Họ và tên",
        rc.get(
            "customer_name",
            "[CHƯA CUNG CẤP]",
        ),
    )

    _add_field(
        doc,
        "Số giấy tờ định danh",
        rc.get(
            "customer_id_number",
            "[CHƯA CUNG CẤP]",
        ),
    )

    _add_field(
        doc,
        "Số tài khoản",
        rc.get(
            "customer_account_number",
            "[CHƯA CUNG CẤP]",
        ),
    )

    _add_field(
        doc,
        "Địa chỉ ví blockchain",
        wallet_from,
    )

    _add_field(
        doc,
        "Ngày sinh",
        state.get(
            "date_of_birth"
        ),
    )

    _add_field(
        doc,
        "Quốc tịch",
        state.get(
            "nationality"
        ),
    )

    _add_field(
        doc,
        "Nghề nghiệp",
        state.get(
            "occupation"
        ),
    )

    _add_field(
        doc,
        "Địa chỉ",
        state.get(
            "customer_address"
        ),
    )

    # ========================================================
    # 8. PHẦN III
    # ========================================================

    _add_section_title(
        doc,
        "PHẦN III. THÔNG TIN VỀ CÁ NHÂN, "
        "TỔ CHỨC CÓ LIÊN QUAN ĐẾN GIAO DỊCH ĐÁNG NGỜ",
    )

    related_parties = (
        state.get(
            "related_parties",
            [],
        )
        or []
    )

    if related_parties:

        rows = []

        for party in related_parties:

            if isinstance(
                party,
                dict,
            ):

                rows.append(
                    [
                        party.get(
                            "name",
                            "Không cung cấp",
                        ),
                        party.get(
                            "type",
                            "Không cung cấp",
                        ),
                        party.get(
                            "relationship",
                            "Không cung cấp",
                        ),
                        party.get(
                            "account",
                            "Không cung cấp",
                        ),
                        party.get(
                            "wallet",
                            "Không cung cấp",
                        ),
                    ]
                )

            else:

                rows.append(
                    [
                        str(party),
                        "Không cung cấp",
                        "Không cung cấp",
                        "Không cung cấp",
                        "Không cung cấp",
                    ]
                )

        _add_table(
            doc,
            [
                "Tên",
                "Loại",
                "Quan hệ",
                "Tài khoản",
                "Ví blockchain",
            ],
            rows,
        )

    else:

        doc.add_paragraph(
            "Chưa xác định cá nhân hoặc tổ chức liên quan "
            "khác tại thời điểm lập báo cáo."
        )

    # ========================================================
    # 9. PHẦN IV
    # ========================================================

    _add_section_title(
        doc,
        "PHẦN IV. THÔNG TIN VỀ GIAO DỊCH ĐÁNG NGỜ",
    )

    _add_field(
        doc,
        "Mã giao dịch",
        tx_hash,
    )

    _add_field(
        doc,
        "Ngày giao dịch",
        state.get(
            "transaction_date"
        ),
    )

    _add_field(
        doc,
        "Loại giao dịch",
        state.get(
            "transaction_type",
            "Giao dịch tài sản số",
        ),
    )

    _add_field(
        doc,
        "Loại tiền",
        state.get(
            "currency",
            "VND equivalent",
        ),
    )

    _add_field(
        doc,
        "Giá trị quy đổi",
        f"{amount_vnd:,.0f} VND",
    )

    _add_field(
        doc,
        "Mục đích giao dịch",
        state.get(
            "transaction_purpose"
        ),
    )

    _add_field(
        doc,
        "Ví gửi",
        wallet_from,
    )

    _add_field(
        doc,
        "Ví nhận",
        wallet_to,
    )

    # --------------------------------------------------------
    # 9.1 Sanctions
    # --------------------------------------------------------

    _add_field(
        doc,
        "OFAC SDN",
        (
            "MATCHED"
            if sanction_result.get(
                "is_match"
            )
            else "NOT MATCHED"
        ),
    )

    if sanction_result.get(
        "is_match"
    ):

        _add_field(
            doc,
            "Ví/thực thể khớp",
            sanction_result.get(
                "matched_wallet",
                "N/A",
            ),
        )

        _add_field(
            doc,
            "Chương trình sanctions",
            sanction_result.get(
                "program",
                "N/A",
            ),
        )

    # --------------------------------------------------------
    # 9.2 Fuzzy name
    # --------------------------------------------------------

    if name_similarity_warning:

        score_text = (
            f"{name_similarity_score:.2f}%"
            if name_similarity_score
            is not None
            else "N/A"
        )

        matched_text = (
            f" — Tên gần khớp: "
            f"{name_similarity_matched_name}"
            if name_similarity_matched_name
            else ""
        )

        _add_field(
            doc,
            "Cảnh báo tương đồng tên",
            (
                f"{score_text}"
                f"{matched_text}. "
                "Đây là fuzzy match, không phải "
                "sanctions match chính xác."
            ),
        )

    # --------------------------------------------------------
    # 9.3 Suspicious transaction description
    # --------------------------------------------------------

    _add_section_title(
        doc,
        "Mô tả và lý do giao dịch được xem xét",
    )

    doc.add_paragraph(
        _build_suspicious_description(
            state
        )
    )

    # ========================================================
    # 10. PHẦN V
    # ========================================================

    _add_section_title(
        doc,
        "PHẦN V. MÔ TẢ CÔNG VIỆC ĐÃ THỰC HIỆN "
        "LIÊN QUAN ĐẾN VIỆC XỬ LÝ BÁO CÁO",
    )

    investigation_steps = (
        state.get(
            "investigation_steps",
            [],
        )
        or []
    )

    if investigation_steps:

        for step in investigation_steps:

            _add_bullet(
                doc,
                step,
            )

    else:

        # Fallback dựa trên những gì hệ thống hiện tại có.
        fallback_steps = [
            "Phân tích giao dịch và thông tin ví blockchain.",
            "Thực hiện sàng lọc sanctions.",
            "Phân tích rủi ro giao dịch.",
        ]

        if community_id is not None:
            fallback_steps.append(
                "Phân tích cấu trúc cộng đồng giao dịch "
                "bằng Graph/Neo4j."
            )

        if suspicious_path:
            fallback_steps.append(
                "Truy vết đường đi đáng ngờ trên đồ thị giao dịch."
            )

        if legal_citations:
            fallback_steps.append(
                "Đối chiếu các căn cứ pháp lý liên quan "
                "thông qua Legal/RAG Agent."
            )

        for step in fallback_steps:

            _add_bullet(
                doc,
                step,
            )

    # ========================================================
    # 11. PHẦN VI
    # ========================================================

    _add_section_title(
        doc,
        "PHẦN VI. CÁC HỒ SƠ, TÀI LIỆU CÓ LIÊN QUAN",
    )

    evidence_documents = (
        state.get(
            "evidence_documents",
            [],
        )
        or []
    )

    if evidence_documents:

        rows = []

        for index, evidence in enumerate(
            evidence_documents,
            start=1,
        ):

            if isinstance(
                evidence,
                dict,
            ):

                rows.append(
                    [
                        index,
                        evidence.get(
                            "type",
                            "Không cung cấp",
                        ),
                        evidence.get(
                            "description",
                            "Không cung cấp",
                        ),
                        evidence.get(
                            "pages",
                            "Không cung cấp",
                        ),
                        evidence.get(
                            "status",
                            "Không cung cấp",
                        ),
                    ]
                )

            else:

                rows.append(
                    [
                        index,
                        "Evidence",
                        str(evidence),
                        "N/A",
                        "Digital",
                    ]
                )

        _add_table(
            doc,
            [
                "STT",
                "Loại hồ sơ",
                "Mô tả",
                "Số trang",
                "Tình trạng",
            ],
            rows,
        )

    else:

        doc.add_paragraph(
            "Chưa có danh mục hồ sơ/tài liệu đính kèm "
            "được cung cấp trong AMLState."
        )

    # ========================================================
    # 12. INTERNAL DECISION SUMMARY
    # ========================================================

    doc.add_page_break()

    _add_section_title(
        doc,
        "PHỤ LỤC A — DIGITALASSET GUARD "
        "INTERNAL RISK ANALYSIS",
    )

    doc.add_paragraph(
        "Phụ lục này là phần phân tích nội bộ của "
        "DigitalAsset Guard nhằm hỗ trợ AML Officer. "
        "Điểm số dưới đây không phải ngưỡng pháp lý "
        "và không tự động quyết định việc gửi STR. "
        "Kể từ khi Decision Engine chuyển sang mô hình rule-based "
        "composite, hệ thống KHÔNG còn tính 1 điểm rủi ro tổng hợp "
        "duy nhất (risk_assessment_score) — quyết định REPORT dựa "
        "trên các bằng chứng độc lập, xem mục Decision Evidence bên "
        "dưới để biết rule cụ thể đã kích hoạt."
    )

    # Risk table — CHỈ hiển thị 2 điểm số nguồn (classifier, graph), KHÔNG
    # còn dòng "Risk Assessment" tổng hợp vì decision_engine.py không tính
    # điểm này nữa (risk_assessment_score luôn = None/0 từ nay — hiển thị
    # 0.0000 ở đây sẽ khiến chuyên viên hiểu nhầm hệ thống đánh giá rủi ro
    # bằng 0, trong khi lý do REPORT thực tế nằm ở Decision Evidence).
    risk_rows = [
        [
            "Classifier",
            f"{float(classifier_score):.4f}",
            "Theo Risk/Classifier Agent"
            + (
                " — KHÔNG dùng làm căn cứ quyết định khi thiếu dữ liệu "
                "on-chain (xem mục Decision Evidence)"
                if state.get("insufficient_data", False)
                else ""
            ),
        ],
        [
            "Graph (PPR)",
            (
                "(không có dữ liệu graph)"
                if state.get("graph_data_available") is not True
                else f"{float(graph_score):.4f}"
            ),
            "Theo Graph Agent — điểm PageRank cá nhân hóa, không phải "
            "căn cứ trực tiếp cho quyết định (xem hop_distance_to_blacklist "
            "trong Graph Evidence bên dưới)",
        ],
        [
            "Đủ dữ liệu on-chain",
            "KHÔNG — ví chưa có lịch sử Etherscan"
            if state.get("insufficient_data", False)
            else "CÓ",
            "FIX 2026-08-08: ví 0 giao dịch (txlist + tokentx) → Decision "
            "Engine bắt buộc route REVIEW thủ công, không REPORT/PASS tự động "
            "dựa trên classifier_score.",
        ],
        [
            "Aggregation / Structuring",
            (
                "ĐÃ ĐÁNH GIÁ — "
                + (
                    "phát hiện dấu hiệu chia nhỏ"
                    if state.get("structuring_detected") is True
                    else "không phát hiện dấu hiệu"
                )
                if state.get("aggregation_status") == "assessed"
                else "CHƯA ĐÁNH GIÁ"
            ),
            (
                "Aggregation Monitor đã chạy rule structuring với dữ liệu "
                "off-chain Core Banking (wallet_tx_history)."
                if state.get("aggregation_status") == "assessed"
                else "Không đánh giá được (thiếu dữ liệu off-chain Core "
                "Banking) — KHÔNG đồng nghĩa đã kiểm tra sạch."
            ),
        ],
    ]

    _add_table(
        doc,
        [
            "Thành phần",
            "Score",
            "Ý nghĩa",
        ],
        risk_rows,
    )

    doc.add_paragraph()

    # --------------------------------------------------------
    # Decision
    # --------------------------------------------------------

    _add_field(
        doc,
        "Decision",
        decision,
    )

    _add_field(
        doc,
        "Case Status",
        case_status,
    )

    _add_field(
        doc,
        "Lý do quyết định",
        decision_reason,
    )

    # --------------------------------------------------------
    # Decision evidence
    # --------------------------------------------------------

    if decision_evidence:

        _add_section_title(
            doc,
            "Decision Evidence",
        )

        for evidence in decision_evidence:

            _add_bullet(
                doc,
                evidence,
            )

    # --------------------------------------------------------
    # Graph evidence
    # --------------------------------------------------------

    if (
        community_id is not None
        or suspicious_path
    ):

        _add_section_title(
            doc,
            "Graph Evidence",
        )

        if community_id is not None:

            _add_field(
                doc,
                "Louvain Community ID",
                community_id,
            )

        if suspicious_path:

            _add_field(
                doc,
                "Suspicious Path",
                " → ".join(
                    map(
                        str,
                        suspicious_path,
                    )
                ),
            )

    # --------------------------------------------------------
    # Sanctions evidence
    # --------------------------------------------------------

    _add_section_title(
        doc,
        "Compliance Evidence",
    )

    _add_field(
        doc,
        "OFAC SDN Match",
        (
            "YES"
            if sanction_result.get(
                "is_match"
            )
            else "NO"
        ),
    )

    if name_similarity_warning:

        _add_field(
            doc,
            "Fuzzy Name Warning",
            "YES",
        )

        _add_field(
            doc,
            "Name Similarity",
            (
                f"{name_similarity_score:.2f}%"
                if name_similarity_score
                is not None
                else "N/A"
            ),
        )

    else:

        _add_field(
            doc,
            "Fuzzy Name Warning",
            "NO",
        )

    # ========================================================
    # 13. LEGAL APPENDIX
    # ========================================================

    if legal_rag_status == "UNAVAILABLE":

        # Thay vì để trống hoặc in nội dung không rõ nguồn gốc, STR phải nêu
        # TƯỜNG MINH rằng căn cứ pháp lý tự động KHÔNG truy xuất được — cảnh báo
        # này bắt buộc chuyên viên AML tự tra cứu trước khi nộp báo cáo.
        _add_section_title(
            doc,
            "PHỤ LỤC B — CĂN CỨ PHÁP LÝ "
            "ĐƯỢC RAG AGENT TRUY XUẤT",
        )

        paragraph = doc.add_paragraph()

        run = paragraph.add_run(
            "⚠️ Không thể truy xuất căn cứ pháp lý tự động "
            "do lỗi hệ thống. Chuyên viên AML phải tự tra cứu "
            "và bổ sung căn cứ pháp lý trước khi nộp báo cáo."
        )

        _set_run_font(
            run,
            bold=True,
        )

        if legal_rag_error:

            error_paragraph = doc.add_paragraph()

            error_run = error_paragraph.add_run(
                f"Chi tiết lỗi hệ thống: "
                f"{legal_rag_error}"
            )

            _set_run_font(
                error_run,
                italic=True,
            )

    elif legal_citations:

        _add_section_title(
            doc,
            "PHỤ LỤC B — CĂN CỨ PHÁP LÝ "
            "ĐƯỢC RAG AGENT TRUY XUẤT",
        )

        _add_legal_citations(
            doc,
            legal_citations,
        )

    # ========================================================
    # 14. FINAL HUMAN REVIEW NOTICE
    # ========================================================

    doc.add_paragraph()

    paragraph = doc.add_paragraph()

    warning = paragraph.add_run(
        "LƯU Ý QUAN TRỌNG: "
        "Đây là DỰ THẢO được hệ thống hỗ trợ tạo. "
        "Báo cáo phải được chuyên viên AML kiểm tra, "
        "bổ sung và phê duyệt trước khi thực hiện "
        "việc gửi báo cáo theo quy trình áp dụng."
    )

    _set_run_font(
        warning,
        bold=True,
    )

    # ========================================================
    # 15. SAVE REPORT
    # ========================================================

    filename = (
        f"STR_REPORT_{tx_hash}.docx"
    )

    report_path = os.path.join(
        REPORTS_OUTPUT_DIR,
        filename,
    )

    doc.save(report_path)

    # ========================================================
    # 16. UPDATE STATE
    # ========================================================

    state["report_path"] = report_path

    # HITL tiếp tục xử lý.
    state["approval_status"] = "pending"

    state["thought"] = (
        f"ReportAssistant: Đã tạo dự thảo STR tại "
        f"{report_path}. Chờ chuyên viên AML duyệt."
    )

    return state


# ============================================================
# LOCAL TEST
# ============================================================

if __name__ == "__main__":

    print(
        "alert_report.py không chạy demo độc lập "
        "vì AMLState phải được tạo từ pipeline chính."
    )

    print(
        "Hãy chạy demo_run.py hoặc pipeline hiện tại "
        "để đảm bảo state schema đồng bộ."
    )